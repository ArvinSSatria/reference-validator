import logging
import os
import io
import docx
import fitz  # PyMuPDF
import xml.etree.ElementTree as ET
from app.utils.text_utils import is_likely_reference

logger = logging.getLogger(__name__)


def get_paragraph_text(paragraph):
    WORD_NAMESPACE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    TEXT = WORD_NAMESPACE + 't'
    xml_string = paragraph._p.xml
    
    try:
        xml_tree = ET.fromstring(xml_string)
        text_nodes = xml_tree.findall('.//' + TEXT)
        if text_nodes:
            return "".join(node.text for node in text_nodes if node.text)
    except ET.ParseError:
        return paragraph.text
    
    return ""


def extract_references_from_docx(file_stream):
    try:
        doc = docx.Document(io.BytesIO(file_stream.read()))
        
        # Gunakan pembaca XML untuk menangani field codes Mendeley
        all_paragraphs = [get_paragraph_text(p).strip() for p in doc.paragraphs]
        paragraphs = [p for p in all_paragraphs if p]  # Filter baris kosong

        REFERENCE_HEADINGS = [
            "daftar pustaka", "daftar referensi", "referensi",
            "bibliography", "references", "pustaka rujukan", "reference"
        ]
        
        start_index = -1
        
        # Cari judul dari depan
        for i, para in enumerate(paragraphs):
            if any(h == para.lower() for h in REFERENCE_HEADINGS):
                # Verifikasi konteks
                for j in range(i + 1, min(i + 6, len(paragraphs))):
                    if is_likely_reference(paragraphs[j]):
                        start_index = j
                        break
                if start_index != -1:
                    break
        
        if start_index == -1:
            return None, "Bagian 'Daftar Pustaka' tidak ditemukan atau tidak diikuti konten valid di file DOCX."

        # PENTING: Deteksi AKHIR section references (seperti teman)
        END_KEYWORDS = [
            'acknowledgments', 'acknowledgements', 'acknowledgment',
            'ucapan terima kasih', 'terima kasih',
            'appendix', 'appendices', 'lampiran',
            'about the authors', 'about authors', 'author information',
            'tentang penulis', 'biography', 'biografi',
            'author contributions', 'kontribusi penulis',
            'funding', 'pendanaan',
            'conflict of interest', 'conflicts of interest', 'konflik kepentingan',
            'competing interests',
            'data availability', 'ketersediaan data',
        ]
        
        end_index = len(paragraphs)  # Default: sampai akhir
        
        # Cari section setelah references
        for i in range(start_index, len(paragraphs)):
            para_lower = paragraphs[i].lower().strip()
            # Cek apakah ini heading (pendek dan match keyword)
            if len(para_lower.split()) <= 5:
                if any(keyword in para_lower for keyword in END_KEYWORDS):
                    end_index = i
                    logger.info(f"✂️ Berhenti di section: '{paragraphs[i]}' (baris #{i})")
                    break
        
        # Ambil paragraf dari start sampai end (BUKAN sampai akhir dokumen!)
        references_block = "\n".join(paragraphs[start_index:end_index])
        
        if not references_block.strip():
            return None, "Bagian 'Daftar Pustaka' ditemukan di DOCX, tetapi isinya kosong."
        
        logger.info(f"✅ Berhasil mengekstrak blok teks referensi dari DOCX")
        logger.info(f"   📄 Total paragraf: {len(paragraphs[start_index:end_index])}")
        logger.info(f"   📏 Panjang teks: {len(references_block)} karakter")
        return references_block, None
        
    except Exception as e:
        logger.error(f"Gagal memproses file DOCX: {e}", exc_info=True)
        # User-friendly error message
        error_msg = "Maaf, terjadi kesalahan saat memproses file DOCX Anda. "
        
        if "password" in str(e).lower() or "encrypted" in str(e).lower():
            error_msg += "File DOCX ter-enkripsi atau dilindungi password. Mohon gunakan file yang tidak terproteksi."
        elif "corrupt" in str(e).lower() or "not a zip file" in str(e).lower():
            error_msg += "File DOCX tampaknya rusak atau tidak valid. Mohon coba file lain."
        elif "memory" in str(e).lower():
            error_msg += "File DOCX terlalu besar untuk diproses. Mohon gunakan file yang lebih kecil."
        else:
            error_msg += "Pastikan file DOCX Anda valid dan dapat dibuka di Microsoft Word."
        
        return None, error_msg


def convert_docx_to_pdf(docx_path):
    """
    Convert DOCX to PDF using PyMuPDF (pure Python, no MS Word needed).
    Returns: (pdf_bytes_stream, error_message)
    """
    try:
        logger.info(f"Converting DOCX to PDF: {docx_path}")
        
        # Read DOCX and extract all text
        doc = docx.Document(docx_path)
        text_content = []
        
        for paragraph in doc.paragraphs:
            text = get_paragraph_text(paragraph).strip()
            if text:
                text_content.append(text)
        
        # Create PDF using PyMuPDF
        pdf_document = fitz.open()  # Create new PDF
        page = pdf_document.new_page(width=595, height=842)  # A4 size
        
        # Write text to PDF
        font_size = 11
        line_height = font_size + 4
        y_position = 50
        margin = 50
        max_width = 595 - (2 * margin)
        
        for text in text_content:
            # Simple text wrapping
            words = text.split()
            line = ""
            
            for word in words:
                test_line = line + word + " "
                # Rough estimate of text width
                if len(test_line) * (font_size * 0.5) > max_width:
                    if line:
                        page.insert_text((margin, y_position), line.strip(), fontsize=font_size)
                        y_position += line_height
                        line = word + " "
                    else:
                        # Word too long, force break
                        page.insert_text((margin, y_position), word, fontsize=font_size)
                        y_position += line_height
                        line = ""
                else:
                    line = test_line
            
            # Write remaining line
            if line:
                page.insert_text((margin, y_position), line.strip(), fontsize=font_size)
                y_position += line_height
            
            # Add extra space after paragraph
            y_position += line_height / 2
            
            # Create new page if needed
            if y_position > 792:  # Near bottom of page
                page = pdf_document.new_page(width=595, height=842)
                y_position = 50
        
        # Save to bytes
        pdf_bytes = pdf_document.tobytes()
        pdf_document.close()
        
        logger.info("DOCX to PDF conversion successful using PyMuPDF")
        return io.BytesIO(pdf_bytes), None
            
    except Exception as e:
        logger.error(f"Error converting DOCX to PDF: {e}", exc_info=True)
        return None, f"Gagal mengonversi DOCX ke PDF: {str(e)}"
